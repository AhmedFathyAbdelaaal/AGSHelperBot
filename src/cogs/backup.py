import discord
from discord import app_commands
from discord.ext import commands
import os
import json
import aiohttp
import asyncio
from docx import Document
from docx.shared import Inches
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import shutil
import datetime
import io
import sqlite3

# Define scopes for Google Drive
SCOPES = ['https://www.googleapis.com/auth/drive']

class Backup(commands.Cog):
    def __init__(self, bot):
        self.bot = bot
        self.drive_service = None
        self.db_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'data', 'backuplogs.db')
        self.ensure_db()
        self.setup_drive_service()

    def ensure_db(self):
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS backup_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                backup_type TEXT,
                target_name TEXT,
                requester_id INTEGER,
                requester_name TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                status TEXT,
                details TEXT
            )
        ''')
        conn.commit()
        conn.close()

    def log_backup(self, backup_type, target_name, requester, status, details):
        try:
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            c.execute('''
                INSERT INTO backup_logs (backup_type, target_name, requester_id, requester_name, status, details)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (backup_type, target_name, requester.id, requester.name, status, details))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Failed to log backup: {e}")

    def setup_drive_service(self):
        """Sets up the Google Drive API service."""
        creds = None
        
        # Option 1: Load from Environment Variable (Best for Coolify/Docker)
        json_creds = os.getenv('GOOGLE_SERVICE_ACCOUNT_JSON')
        if json_creds:
            try:
                info = json.loads(json_creds)
                creds = service_account.Credentials.from_service_account_info(info, scopes=SCOPES)
                print("Google Drive Service initialized from Environment Variable.")
            except Exception as e:
                print(f"Failed to load credentials from Environment Variable: {e}")

        # Option 2: Load from File (Best for Local Dev)
        if not creds:
            creds_path = 'service_account.json'
            if os.path.exists(creds_path):
                try:
                    creds = service_account.Credentials.from_service_account_file(
                        creds_path, scopes=SCOPES)
                    print("Google Drive Service initialized from file.")
                except Exception as e:
                    print(f"Failed to initialize Google Drive Service from file: {e}")
            else:
                print("No credentials found (checked ENV and service_account.json). Google Drive backup will not work.")
                return

        if creds:
            try:
                self.drive_service = build('drive', 'v3', credentials=creds)
            except Exception as e:
                print(f"Failed to build Drive service: {e}")

    # --- Google Drive Helper Functions ---

    async def create_drive_folder(self, folder_name, parent_id=None):
        """Creates a folder in Google Drive (Async wrapper)."""
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, self._create_drive_folder_sync, folder_name, parent_id)

    def _create_drive_folder_sync(self, folder_name, parent_id=None):
        """Creates a folder in Google Drive (Blocking)."""
        if not self.drive_service:
            return None
        
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            file_metadata['parents'] = [parent_id]
        
        try:
            file = self.drive_service.files().create(body=file_metadata, fields='id', supportsAllDrives=True).execute()
            return file.get('id')
        except Exception as e:
            print(f"Error creating folder {folder_name}: {e}")
            return None

    async def upload_file_to_drive(self, file_path, folder_id=None):
        """Uploads a file to Google Drive (Async wrapper)."""
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, self._upload_file_to_drive_sync, file_path, folder_id)

    def _upload_file_to_drive_sync(self, file_path, folder_id=None):
        """Uploads a file to Google Drive (Blocking)."""
        if not self.drive_service:
            print("Drive service not initialized.")
            return None
        
        if not os.path.exists(file_path):
            print(f"File not found locally: {file_path}")
            return None

        file_name = os.path.basename(file_path)
        print(f"Uploading {file_name} to Drive folder ID {folder_id}...")
        
        file_metadata = {'name': file_name}
        if folder_id:
            file_metadata['parents'] = [folder_id]
        
        try:
            media = MediaFileUpload(file_path, resumable=True)
            file = self.drive_service.files().create(body=file_metadata, media_body=media, fields='id', supportsAllDrives=True).execute()
            print(f"Successfully uploaded {file_name} (ID: {file.get('id')})")
            return file.get('id')
        except Exception as e:
            print(f"Error uploading file {file_name}: {e}")
            return None

    async def share_file_anyone_reader(self, file_id):
        """Makes a file or folder accessible to anyone with the link (Async wrapper)."""
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(None, self._share_file_anyone_reader_sync, file_id)

    def _share_file_anyone_reader_sync(self, file_id):
        """Makes a file or folder accessible to anyone with the link (Blocking)."""
        if not self.drive_service:
            return None
        
        user_permission = {
            'type': 'anyone',
            'role': 'reader',
        }
        try:
            self.drive_service.permissions().create(
                fileId=file_id,
                body=user_permission,
                fields='id',
                supportsAllDrives=True
            ).execute()
            print(f"Shared file {file_id} with anyone.")
        except Exception as e:
            print(f"Error sharing file {file_id}: {e}")

    # --- Local Helper Functions ---

    async def download_attachment(self, url, save_path):
        """Downloads an attachment from a URL."""
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as resp:
                if resp.status == 200:
                    with open(save_path, 'wb') as f:
                        f.write(await resp.read())
                    return True
        return False

    def clean_filename(self, filename):
        """Removes illegal characters from filenames."""
        cleaned = "".join([c for c in filename if c.isalpha() or c.isdigit() or c in (' ', '.', '_', '-')]).rstrip()
        if not cleaned:
            return "unnamed_channel"
        return cleaned

    async def create_docx_from_messages(self, messages, output_path, attachments_dir):
        """Creates a docx file from a list of messages."""
        print(f"Creating docx with {len(messages)} messages at {output_path}")
        doc = Document()
        doc.add_heading('Channel Backup', 0)

        for msg in reversed(messages): # Messages are usually fetched newest first
            p = doc.add_paragraph()
            timestamp = msg.created_at.strftime("%Y-%m-%d %H:%M:%S")
            runner = p.add_run(f"[{timestamp}] {msg.author.name}: ")
            runner.bold = True
            p.add_run(msg.content)

            if msg.attachments:
                for attachment in msg.attachments:
                    file_name = self.clean_filename(attachment.filename)
                    save_path = os.path.join(attachments_dir, file_name)
                    
                    # Ensure attachments dir exists
                    os.makedirs(attachments_dir, exist_ok=True)

                    if await self.download_attachment(attachment.url, save_path):
                        # Check if image and embed
                        if attachment.content_type and attachment.content_type.startswith('image/'):
                            try:
                                doc.add_picture(save_path, width=Inches(4.0))
                            except Exception:
                                doc.add_paragraph(f"[Could not embed image: {file_name}]")
                        else:
                            doc.add_paragraph(f"[Attachment: {file_name}]")
                    else:
                        doc.add_paragraph(f"[Failed to download attachment: {file_name}]")
        
        doc.save(output_path)

    # --- Backup Logic ---

    async def backup_thread_logic(self, thread, parent_drive_id, temp_dir, after_date=None):
        """Backs up a single thread."""
        print(f"[LOG] Starting backup for thread: {thread.name} (ID: {thread.id})")
        thread_name = self.clean_filename(thread.name)
        thread_folder_local = os.path.join(temp_dir, thread_name)
        os.makedirs(thread_folder_local, exist_ok=True)
        
        attachments_folder_local = os.path.join(thread_folder_local, "Attachments")
        os.makedirs(attachments_folder_local, exist_ok=True)

        docx_path = os.path.join(thread_folder_local, f"{thread_name}.docx")

        print(f"[LOG] Fetching messages for thread: {thread.name}...")
        messages = [message async for message in thread.history(limit=None, after=after_date)]
        print(f"[LOG] Fetched {len(messages)} messages for thread: {thread.name}. Creating DOCX...")
        await self.create_docx_from_messages(messages, docx_path, attachments_folder_local)

        # Upload to Drive
        # 1. Create Thread Folder in Drive
        print(f"[LOG] Creating Drive folder for thread: {thread.name}...")
        thread_drive_id = await self.create_drive_folder(thread_name, parent_drive_id)
        
        # 2. Upload Docx
        print(f"[LOG] Uploading DOCX for thread: {thread.name}...")
        await self.upload_file_to_drive(docx_path, thread_drive_id)

        # 3. Upload Attachments Folder
        if os.listdir(attachments_folder_local):
            print(f"[LOG] Uploading attachments for thread: {thread.name}...")
            attachments_drive_id = await self.create_drive_folder("Attachments", thread_drive_id)
            for filename in os.listdir(attachments_folder_local):
                file_path = os.path.join(attachments_folder_local, filename)
                await self.upload_file_to_drive(file_path, attachments_drive_id)
        
        print(f"[LOG] Finished backup for thread: {thread.name}")
        return thread_drive_id

    async def backup_channel_logic(self, channel, parent_drive_id, temp_dir, after_date=None):
        """Backs up a channel and its threads."""
        print(f"[LOG] Starting backup for channel: {channel.name} (ID: {channel.id})")
        channel_name = self.clean_filename(channel.name)
        channel_folder_local = os.path.join(temp_dir, channel_name)
        os.makedirs(channel_folder_local, exist_ok=True)

        attachments_folder_local = os.path.join(channel_folder_local, "Attachments")
        os.makedirs(attachments_folder_local, exist_ok=True)
        
        threads_folder_local = os.path.join(channel_folder_local, "Threads")
        # We don't makedirs threads yet, only if needed

        docx_path = os.path.join(channel_folder_local, f"{channel_name}.docx")

        # Fetch messages
        print(f"[LOG] Fetching messages for channel: {channel.name}...")
        messages = [message async for message in channel.history(limit=None, after=after_date)]
        print(f"[LOG] Fetched {len(messages)} messages for channel: {channel.name}. Creating DOCX...")
        await self.create_docx_from_messages(messages, docx_path, attachments_folder_local)

        # Create Channel Folder in Drive
        print(f"[LOG] Creating Drive folder for channel: {channel.name}...")
        channel_drive_id = await self.create_drive_folder(channel_name, parent_drive_id)

        # Upload Docx
        print(f"[LOG] Uploading DOCX for channel: {channel.name}...")
        await self.upload_file_to_drive(docx_path, channel_drive_id)

        # Upload Attachments
        if os.listdir(attachments_folder_local):
            print(f"[LOG] Uploading attachments for channel: {channel.name}...")
            attachments_drive_id = await self.create_drive_folder("Attachments", channel_drive_id)
            for filename in os.listdir(attachments_folder_local):
                await self.upload_file_to_drive(os.path.join(attachments_folder_local, filename), attachments_drive_id)

        # Handle Threads
        if hasattr(channel, 'threads'):
            # Archived threads need to be fetched separately usually, but active are in .threads
            # For simplicity, we iterate active threads. 
            # To get archived, we'd need channel.archived_threads()
            
            all_threads = channel.threads
            # async for t in channel.archived_threads(): all_threads.append(t) # Optional: Add archived

            if all_threads:
                print(f"[LOG] Found {len(all_threads)} threads in channel: {channel.name}. Backing them up...")
                threads_drive_id = await self.create_drive_folder("Threads", channel_drive_id)
                os.makedirs(threads_folder_local, exist_ok=True)
                
                for thread in all_threads:
                    await self.backup_thread_logic(thread, threads_drive_id, threads_folder_local, after_date)
        
        print(f"[LOG] Finished backup for channel: {channel.name}")
        return channel_drive_id

    async def backup_category_logic(self, category, parent_drive_id, temp_dir, after_date=None):
        """Backs up a category and its channels."""
        print(f"[LOG] Starting backup for category: {category.name} (ID: {category.id})")
        category_name = self.clean_filename(category.name)
        category_folder_local = os.path.join(temp_dir, category_name)
        os.makedirs(category_folder_local, exist_ok=True)

        # Create Category Folder in Drive
        print(f"[LOG] Creating Drive folder for category: {category.name}...")
        category_drive_id = await self.create_drive_folder(category_name, parent_drive_id)

        print(f"[LOG] Processing {len(category.channels)} channels in category: {category.name}...")
        for channel in category.channels:
            if isinstance(channel, discord.TextChannel):
                await self.backup_channel_logic(channel, category_drive_id, category_folder_local, after_date)
        
        print(f"[LOG] Finished backup for category: {category.name}")
        return category_drive_id

    # --- Helper Methods ---

    async def check_perms(self, interaction: discord.Interaction) -> bool:
        """Checks if the user has the required role."""
        # Block specific user
        if interaction.user.id == 327523870464540673: # REPLACE THIS WITH THE USER ID
            await interaction.response.send_message("🚫 You are blocked from using this bot.", ephemeral=True)
            return False

        if not any(role.name == "hasBotPerms" for role in interaction.user.roles):
            await interaction.response.send_message("🚫 You are not allowed to use this bot. You need the 'hasBotPerms' role.", ephemeral=True)
            return False
        return True

    # --- Commands ---

    backup_group = app_commands.Group(name="backup", description="Backup commands")

    def parse_duration(self, duration_str: str):
        """Parses a duration string (e.g., '7d', '24h') into a datetime object."""
        if not duration_str:
            return None
        
        now = datetime.datetime.now(datetime.timezone.utc)
        try:
            amount = int(duration_str[:-1])
            unit = duration_str[-1].lower()
            
            if unit == 'd':
                delta = datetime.timedelta(days=amount)
            elif unit == 'h':
                delta = datetime.timedelta(hours=amount)
            elif unit == 'w':
                delta = datetime.timedelta(weeks=amount)
            elif unit == 'm': # Minutes
                delta = datetime.timedelta(minutes=amount)
            else:
                return None
            
            return now - delta
        except ValueError:
            return None

    @backup_group.command(name="thread", description="Back up a thread")
    @app_commands.describe(since="Backup messages since (e.g., 7d, 24h, 2w)")
    async def backup_thread(self, interaction: discord.Interaction, thread: discord.Thread = None, since: str = None):
        """Backs up a thread."""
        if not await self.check_perms(interaction):
            return

        target_thread = thread or interaction.channel
        if not isinstance(target_thread, discord.Thread):
            await interaction.response.send_message("The target is not a thread.", ephemeral=True)
            return

        after_date = self.parse_duration(since)
        if since and not after_date:
            await interaction.response.send_message("Invalid duration format. Use 7d, 24h, 2w, etc.", ephemeral=True)
            return

        if not self.drive_service:
            await interaction.response.send_message("Google Drive service is not configured.", ephemeral=True)
            return

        await interaction.response.defer(ephemeral=True, thinking=True)
        
        root_folder_id = os.getenv('BACKUP_ROOT_FOLDER_ID')
        if not root_folder_id:
            await interaction.followup.send("BACKUP_ROOT_FOLDER_ID not found in .env")
            return

        temp_dir = f"temp_backup_{interaction.id}"
        try:
            threads_root_id = await self.create_drive_folder("Threads Folder", root_folder_id)
            drive_id = await self.backup_thread_logic(target_thread, threads_root_id, temp_dir, after_date)
            await self.share_file_anyone_reader(drive_id)
            link = f"https://drive.google.com/drive/folders/{drive_id}"
            await interaction.followup.send(f"Thread backup complete! [View in Drive]({link})")
        except Exception as e:
            await interaction.followup.send(f"Backup failed: {e}")
            print(e)
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

    @backup_group.command(name="channel", description="Back up a channel")
    @app_commands.describe(since="Backup messages since (e.g., 7d, 24h, 2w)")
    async def backup_channel(self, interaction: discord.Interaction, channel: discord.TextChannel = None, since: str = None):
        """Backs up a channel."""
        if not await self.check_perms(interaction):
            return

        target_channel = channel or interaction.channel
        if not isinstance(target_channel, discord.TextChannel):
             await interaction.response.send_message("The target is not a text channel.", ephemeral=True)
             return

        after_date = self.parse_duration(since)
        if since and not after_date:
            await interaction.response.send_message("Invalid duration format. Use 7d, 24h, 2w, etc.", ephemeral=True)
            return

        if not self.drive_service:
            await interaction.response.send_message("Google Drive service is not configured.", ephemeral=True)
            return

        await interaction.response.defer(ephemeral=True, thinking=True)
        
        root_folder_id = os.getenv('BACKUP_ROOT_FOLDER_ID')
        if not root_folder_id:
            await interaction.followup.send("BACKUP_ROOT_FOLDER_ID not found in .env")
            return

        temp_dir = f"temp_backup_{interaction.id}"
        try:
            drive_id = await self.backup_channel_logic(target_channel, root_folder_id, temp_dir, after_date)
            await self.share_file_anyone_reader(drive_id)
            link = f"https://drive.google.com/drive/folders/{drive_id}"
            await interaction.followup.send(f"Channel backup complete! [View in Drive]({link})")
            self.log_backup('channel', target_channel.name, interaction.user, 'Success', link)
        except Exception as e:
            self.log_backup('channel', target_channel.name, interaction.user, 'Failed', str(e))
            await interaction.followup.send(f"Backup failed: {e}")
            print(e)
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

    @backup_group.command(name="category", description="Back up a category")
    @app_commands.describe(since="Backup messages since (e.g., 7d, 24h, 2w)")
    async def backup_category(self, interaction: discord.Interaction, category: discord.CategoryChannel = None, since: str = None):
        """Backs up a category."""
        if not await self.check_perms(interaction):
            return

        target_category = category or interaction.channel.category
        if not target_category:
            await interaction.response.send_message("No category specified or found.", ephemeral=True)
            return

        after_date = self.parse_duration(since)
        if since and not after_date:
            await interaction.response.send_message("Invalid duration format. Use 7d, 24h, 2w, etc.", ephemeral=True)
            return

        if not self.drive_service:
            await interaction.response.send_message("Google Drive service is not configured.", ephemeral=True)
            return

        await interaction.response.defer(ephemeral=True, thinking=True)
        
        root_folder_id = os.getenv('BACKUP_ROOT_FOLDER_ID')
        if not root_folder_id:
            await interaction.followup.send("BACKUP_ROOT_FOLDER_ID not found in .env")
            return

        temp_dir = f"temp_backup_{interaction.id}"
        try:
            drive_id = await self.backup_category_logic(target_category, root_folder_id, temp_dir, after_date)
            await self.share_file_anyone_reader(drive_id)
            link = f"https://drive.google.com/drive/folders/{drive_id}"
            await interaction.followup.send(f"Category backup complete! [View in Drive]({link})")
            self.log_backup('category', target_category.name, interaction.user, 'Success', link)
        except Exception as e:
            self.log_backup('category', target_category.name, interaction.user, 'Failed', str(e))
            await interaction.followup.send(f"Backup failed: {e}")
            print(e)
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

    @backup_group.command(name="server", description="Back up the entire server")
    @app_commands.describe(since="Backup messages since (e.g., 7d, 24h, 2w)")
    async def backup_server(self, interaction: discord.Interaction, since: str = None):
        """Backs up the entire server."""
        if not await self.check_perms(interaction):
            return

        # Check for isAdmin role
        if not any(role.name == "isAdmin" for role in interaction.user.roles):
            await interaction.response.send_message("🚫 You need the 'isAdmin' role to use this command.", ephemeral=True)
            return

        after_date = self.parse_duration(since)
        if since and not after_date:
            await interaction.response.send_message("Invalid duration format. Use 7d, 24h, 2w, etc.", ephemeral=True)
            return

        if not self.drive_service:
            await interaction.response.send_message("Google Drive service is not configured.", ephemeral=True)
            return

        await interaction.response.defer(ephemeral=True, thinking=True)
        
        root_folder_id = os.getenv('BACKUP_ROOT_FOLDER_ID')
        if not root_folder_id:
            await interaction.followup.send("BACKUP_ROOT_FOLDER_ID not found in .env")
            return

        temp_dir = f"temp_backup_{interaction.id}"
        try:
            print(f"[LOG] Starting SERVER backup for: {interaction.guild.name}")
            # Folder: Server Name
            server_name = self.clean_filename(interaction.guild.name)
            server_drive_id = await self.create_drive_folder(server_name, root_folder_id)
            
            # Categories
            print(f"[LOG] Processing {len(interaction.guild.categories)} categories...")
            for category in interaction.guild.categories:
                await self.backup_category_logic(category, server_drive_id, temp_dir, after_date)
            
            # Channels not in categories?
            # Discord channels can exist without categories.
            no_cat_channels = [c for c in interaction.guild.text_channels if not c.category]
            if no_cat_channels:
                print(f"[LOG] Processing {len(no_cat_channels)} uncategorized channels...")
                uncategorized_id = await self.create_drive_folder("Uncategorized", server_drive_id)
                uncategorized_dir = os.path.join(temp_dir, "Uncategorized")
                for channel in no_cat_channels:
                    await self.backup_channel_logic(channel, uncategorized_id, uncategorized_dir, after_date)

            await self.share_file_anyone_reader(server_drive_id)
            link = f"https://drive.google.com/drive/folders/{server_drive_id}"
            print(f"[LOG] Server backup complete for: {interaction.guild.name}")
            await interaction.followup.send(f"Server backup complete! [View in Drive]({link})")
            self.log_backup('server', interaction.guild.name, interaction.user, 'Success', link)
        except Exception as e:
            self.log_backup('server', interaction.guild.name, interaction.user, 'Failed', str(e))
            await interaction.followup.send(f"Backup failed: {e}")
            print(f"[ERROR] Backup failed: {e}")
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

async def setup(bot):
    await bot.add_cog(Backup(bot))
